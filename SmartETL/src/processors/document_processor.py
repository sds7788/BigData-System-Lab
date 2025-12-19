"""
多格式文档处理器 - 支持PDF和DOCX
"""
import os
import hashlib
import re
import sys
import io
from typing import Optional, Dict

class DocumentProcessor:
    def __init__(self, file_path: str, original_url: Optional[str] = None):
        self.file_path = file_path
        self.original_url = original_url
        self.doc_id = hashlib.md5(str(file_path).encode()).hexdigest()
        self.file_ext = os.path.splitext(file_path)[1].lower()

    def process(self) -> Optional[Dict]:
        """处理文档文件的完整流程"""
        try:
            # 只处理PDF和DOCX文件
            if self.file_ext not in ['.pdf', '.docx']:
                print(f"  跳过不支持的文件格式: {self.file_ext}")
                return None

            full_text = self.extract_text()
            if not full_text or len(full_text.strip()) < 10:
                return None

            metadata = self.extract_metadata(full_text)

            if ('表' in metadata['title']):
                return None

            return {
                'id': self.doc_id,
                'title': metadata['title'],
                'year': metadata['year'],
                'department': metadata['department'],
                'full_text': full_text,
                'file_path': str(self.file_path),
                'file_type': self.file_ext[1:]
            }

        except Exception as e:
            print(f"处理文档 {os.path.basename(self.file_path)} 时出错: {e}")
            return None

    def extract_text(self) -> str:
        if self.file_ext == '.pdf':
            return self._extract_pdf_text()
        elif self.file_ext == '.docx':
            return self._extract_docx_text()
        else:
            return ""

    def _extract_pdf_text(self) -> str:
        """提取PDF文本"""
        try:
            import pdfplumber

            old_stderr = sys.stderr
            sys.stderr = io.StringIO()

            full_text = ""
            with pdfplumber.open(self.file_path) as pdf:
                for page in pdf.pages:
                    try:
                        text = page.extract_text()
                        if text:
                            full_text += text + "\n"
                    except:
                        continue

            sys.stderr = old_stderr
            return full_text.strip()

        except ImportError:
            print("  需要安装pdfplumber: pip install pdfplumber")
            return ""
        except Exception as e:
            print(f"  提取PDF文本失败: {e}")
            return ""

    def _extract_docx_text(self) -> str:
        """提取DOCX文本"""
        try:
            import docx

            doc = docx.Document(self.file_path)
            full_text = []

            # 提取段落文本
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    full_text.append(text)

            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text:
                            row_text.append(text)
                    if row_text:
                        full_text.append(" | ".join(row_text))

            return "\n".join(full_text)

        except ImportError:
            print("  需要安装python-docx: pip install python-docx")
            return ""
        except Exception as e:
            print(f"  提取DOCX文本失败: {e}")
            return ""

    def extract_metadata(self, text: str) -> Dict:
        """从文本中提取元数据"""
        return {
            'title': self._extract_title(text),
            'year': self._extract_year(text),
            'department': self._extract_department()
        }

    def _extract_title(self, text: str) -> str:
        """提取标题 - 取前两行文本"""
        if not text:
            filename = os.path.basename(self.file_path)
            return filename.replace(self.file_ext, '').replace('_', ' ').replace('-', ' ')

        lines = [line.strip() for line in text.split('\n') if line.strip()]

        if len(lines) >= 2:
            title = f"{lines[0]} {lines[1]}"
            title = title.replace('\\n', ' ').replace('\r', ' ')[:150]
            return title
        elif len(lines) == 1:
            title = lines[0].replace('\\n', ' ').replace('\r', ' ')[:100]
            return title
        else:
            # 没有文本内容，使用文件名
            filename = os.path.basename(self.file_path)
            return filename.replace(self.file_ext, '')

    def _extract_year(self, text: str) -> str:
        """提取年份"""
        if not text:
            return "未知"

        year_patterns = [
            r'(\d{4})年',
            r'(\d{4})年度',
            r'[\(（](\d{4})[\)）]',
            r'\b(20\d{2})\b',
            r'\b(19\d{2})\b'
        ]

        for pattern in year_patterns:
            matches = re.findall(pattern, text)
            if matches:
                year = matches[0] if isinstance(matches[0], str) else matches[0][0]
                if 1900 <= int(year) <= 2100:
                    return year

        filename = os.path.basename(self.file_path)
        year_match = re.search(r'(\d{4})', filename)
        if year_match:
            year = year_match.group(1)
            if 1900 <= int(year) <= 2100:
                return year

        return "2025"

    def _extract_department(self) -> str:
        """根据URL提取发布部门"""
        if not self.original_url:
            return "其他"

        url_lower = self.original_url.lower()

        if 'finance' in url_lower:
            return '财务处'
        elif 'teach' in url_lower:
            return '教务处'
        elif 'saids' in url_lower:
            return '人工智能与数据科学学院'
        elif 'cs' in url_lower:
            return '计算机科学与技术学院'
        elif 'sist' in url_lower:
            return '信息科学技术学院'
        elif 'zsb' in url_lower:
            return '本科生招生网'
        elif 'job' in url_lower:
            return '就业信息网'

        return "其他"
